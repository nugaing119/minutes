from __future__ import annotations

import argparse
import io
import os
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor, Twips


BLUE_DARK = "1F3864"
BLUE = "2E75B6"
BLUE_LIGHT = "D5E8F0"
GRAY_LIGHT = "F2F2F2"
ORANGE_LIGHT = "FFF3E0"
BORDER = "CCCCCC"
WHITE = "FFFFFF"
DOCUMENTS_PRESET = "retained_word_template"
HEADER_PATTERN = "heatwave_reference_cover"
WORD_TEMPLATE_ID = "heatwave-reference-v1"
DEFAULT_WORD_TEMPLATE_PATH = (
    Path(__file__).resolve().parents[1]
    / "codex"
    / "skills"
    / "minutes"
    / "assets"
    / "minutes-word-template.docx"
)
BODY_FONT = "Arial"
CODE_FONT = "Courier New"
TOC_DETAIL_ENTRY_LIMIT = 24
CHECKBOX_ITEM_PATTERN = re.compile(r"^\[(?P<mark>[ xX])\]\s+(?P<text>.+)$")
INLINE_MARKDOWN_PATTERN = re.compile(
    r"`(?P<code>[^`]+)`"
    r"|\[(?P<link_text>[^\]]+)\]\((?P<link_url>https?://[^)]+)\)"
    r"|\*\*(?P<strong>.+?)\*\*"
    r"|(?<!\*)\*(?P<emphasis>[^*]+?)\*(?!\*)"
)
MARKDOWN_IMAGE_PATTERN = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")
HTML_BREAK_PATTERN = re.compile(r"<br\s*/?>", re.I)
DANGLING_EVIDENCE_PATTERN = re.compile(
    r"\s+Evidence:\s*(?:and\s*)?\.(?=(?:\*)?$)",
    re.I,
)
MANUAL_HEADING_PREFIX_PATTERN = re.compile(
    r"^(?:\d+[.)]|\d+(?:\.\d+)+)\s+"
)
JFIF_APP0_SEGMENT = (
    b"\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00"
)
INTERNAL_CITATION_PATTERN = re.compile(
    r":codex-file-citation\{|cite|\bturn\d+(?:search|fetch|view)\d+\b",
    re.I,
)
DOCUMENT_UI_LABELS = {
    "en": {"contents": "Contents", "recorded": "Recorded"},
    "ko": {"contents": "목차", "recorded": "작성일"},
}


@dataclass
class MarkdownTable:
    headers: list[str]
    rows: list[list[str]]


@dataclass
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    items: list[str] | None = None
    table: MarkdownTable | None = None
    path: str = ""


def generate_docx_report(
    markdown_path: Path,
    output_path: Path,
    document_title: str | None = None,
    saved_date: str = "",
    document_type: str | None = None,
    meeting_title: str | None = None,
    template_path: Path | None = None,
) -> Path:
    markdown = markdown_path.read_text(encoding="utf-8")
    front_matter, document_markdown = split_front_matter(markdown)
    validate_supported_markdown(document_markdown)
    blocks = parse_markdown(document_markdown)
    title_index = first_document_title_index(blocks)
    markdown_title = blocks[title_index].text if title_index is not None else ""
    resolved_title = (
        (document_title or "").strip()
        or (meeting_title or "").strip()
        or front_matter.get("title", "").strip()
        or markdown_title.strip()
        or output_path.stem
    )
    resolved_type = (
        (document_type or "").strip()
        or front_matter.get("document_type", "").strip()
        or extract_document_type(document_markdown)
    )
    document_language = detect_document_language(document_markdown)
    heading_anchors = build_heading_anchors(blocks, title_index=title_index)
    heading_numbers, numbering_base_level = build_heading_numbers(
        blocks,
        title_index=title_index,
    )
    use_automatic_heading_numbers = bool(heading_numbers) and not has_manual_heading_numbers(
        blocks,
        heading_numbers,
    )

    doc = load_word_template(template_path)
    heading_num_id = (
        add_heading_numbering_definition(doc)
        if use_automatic_heading_numbers
        else None
    )
    add_cover(
        doc,
        resolved_title,
        saved_date,
        resolved_type,
        document_language=document_language,
    )
    add_static_toc(
        doc,
        blocks,
        heading_anchors,
        heading_numbers=heading_numbers if use_automatic_heading_numbers else {},
        title_index=title_index,
        document_language=document_language,
    )
    add_body(
        doc,
        blocks,
        heading_anchors,
        heading_num_id=heading_num_id,
        numbering_base_level=numbering_base_level,
        numbered_heading_indexes=set(heading_numbers) if use_automatic_heading_numbers else set(),
        title_index=title_index,
        source_dir=markdown_path.parent,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def resolve_word_template_path(template_path: Path | None = None) -> Path:
    configured = template_path or Path(
        os.environ.get("MINUTES_WORD_TEMPLATE", str(DEFAULT_WORD_TEMPLATE_PATH))
    )
    resolved = configured.expanduser().resolve()
    if not resolved.is_file():
        raise FileNotFoundError(
            f"minutes Word template is missing: {resolved}. "
            "Restore codex/skills/minutes/assets/minutes-word-template.docx."
        )
    return resolved


def load_word_template(template_path: Path | None = None) -> Document:
    """Load the retained template and clear only its editable body slot."""
    document = Document(resolve_word_template_path(template_path))
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return document


def split_front_matter(markdown: str) -> tuple[dict[str, str], str]:
    lines = markdown.splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, markdown
    closing_index = next(
        (index for index, line in enumerate(lines[1:], start=1) if line.strip() == "---"),
        None,
    )
    if closing_index is None:
        raise ValueError("unterminated Markdown YAML front matter")
    metadata: dict[str, str] = {}
    for line in lines[1:closing_index]:
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        key, separator, value = stripped.partition(":")
        if not separator or not key.strip():
            raise ValueError(f"unsupported Markdown front matter line: {line}")
        clean_value = value.strip()
        if (
            len(clean_value) >= 2
            and clean_value[0] == clean_value[-1]
            and clean_value[0] in {'"', "'"}
        ):
            clean_value = clean_value[1:-1]
        metadata[key.strip()] = clean_value
    body = "\n".join(lines[closing_index + 1 :])
    if markdown.endswith("\n"):
        body += "\n"
    return metadata, body


def parse_markdown(markdown: str) -> list[MarkdownBlock]:
    lines = markdown.splitlines()
    blocks: list[MarkdownBlock] = []
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            blocks.append(
                MarkdownBlock(
                    kind="heading",
                    level=len(heading.group(1)),
                    text=heading.group(2).strip(),
                )
            )
            index += 1
            continue

        image = MARKDOWN_IMAGE_PATTERN.fullmatch(line.strip())
        if image:
            blocks.append(
                MarkdownBlock(
                    kind="image",
                    text=image.group("alt").strip(),
                    path=image.group("path").strip(),
                )
            )
            index += 1
            continue

        if is_table_start(lines, index):
            table, index = parse_table(lines, index)
            blocks.append(MarkdownBlock(kind="table", table=table))
            continue

        if line.lstrip().startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].lstrip().startswith(">"):
                quote_lines.append(lines[index].lstrip()[1:].lstrip())
                index += 1
            blocks.append(MarkdownBlock(kind="blockquote", text=" ".join(quote_lines)))
            continue

        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while index < len(lines) and re.match(r"^\s*[-*]\s+", lines[index]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[index]).strip())
                index += 1
            blocks.append(MarkdownBlock(kind="bullets", items=items))
            continue

        if re.match(r"^\s*\d+[.)]\s+", line):
            items = []
            while index < len(lines) and re.match(
                r"^\s*\d+[.)]\s+",
                lines[index],
            ):
                items.append(
                    re.sub(r"^\s*\d+[.)]\s+", "", lines[index]).strip()
                )
                index += 1
            blocks.append(MarkdownBlock(kind="numbered", items=items))
            continue

        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines):
            next_line = lines[index].rstrip()
            if not next_line.strip():
                break
            if re.match(r"^(#{1,6})\s+(.+)$", next_line):
                break
            if (
                re.match(r"^\s*[-*]\s+", next_line)
                or re.match(r"^\s*\d+[.)]\s+", next_line)
                or next_line.lstrip().startswith(">")
                or MARKDOWN_IMAGE_PATTERN.fullmatch(next_line.strip())
                or is_table_start(lines, index)
            ):
                break
            paragraph_lines.append(next_line.strip())
            index += 1
        blocks.append(MarkdownBlock(kind="paragraph", text=" ".join(paragraph_lines)))
    return blocks


def validate_supported_markdown(markdown: str) -> None:
    if INTERNAL_CITATION_PATTERN.search(markdown):
        raise ValueError("internal Codex citation tokens are not allowed in DOCX source")
    if markdown.count("**") % 2:
        raise ValueError("unbalanced bold Markdown marker: **")
    if markdown.count("`") % 2:
        raise ValueError("unbalanced inline code Markdown marker: `")


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return lines[index].strip().startswith("|") and re.match(
        r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$",
        lines[index + 1].strip(),
    )


def parse_table(lines: list[str], index: int) -> tuple[MarkdownTable, int]:
    headers = split_table_row(lines[index])
    index += 2
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append(split_table_row(lines[index]))
        index += 1
    width = len(headers)
    rows = [row + [""] * (width - len(row)) for row in rows]
    return MarkdownTable(headers=headers, rows=[row[:width] for row in rows]), index


def split_table_row(line: str) -> list[str]:
    clean = line.strip().strip("|")
    return [cell.replace("\\|", "|").strip() for cell in clean.split("|")]


def first_document_title_index(blocks: list[MarkdownBlock]) -> int | None:
    for index, block in enumerate(blocks):
        if block.kind == "heading" and block.level == 1:
            return index
    return None


def extract_document_type(markdown: str) -> str:
    for line in markdown.splitlines():
        match = re.match(
            r"^(?:[-*]\s*)?(?:문서 유형|Document type)\s*:\s*(.+)$",
            line.strip(),
            re.I,
        )
        if match:
            return match.group(1).strip()
    return ""


def detect_document_language(markdown: str) -> str:
    if re.search(r"(?im)^\s*(?:[-*]\s*)?Document type\s*:", markdown):
        return "en"
    if re.search(r"(?m)^\s*(?:[-*]\s*)?문서 유형\s*:", markdown):
        return "ko"
    return "ko" if re.search(r"[가-힣]", markdown) else "en"


def build_heading_anchors(
    blocks: list[MarkdownBlock],
    *,
    title_index: int | None = None,
) -> dict[int, str]:
    anchors = {}
    count = 1
    for index, block in enumerate(blocks):
        if block.kind != "heading" or block.level > 3:
            continue
        if index == title_index:
            continue
        anchors[index] = f"meeting_section_{count:04d}"
        count += 1
    return anchors


def build_heading_numbers(
    blocks: list[MarkdownBlock],
    *,
    title_index: int | None = None,
) -> tuple[dict[int, str], int | None]:
    headings = [
        (index, block)
        for index, block in enumerate(blocks)
        if block.kind == "heading"
        and index != title_index
        and block.level <= 3
    ]
    if not headings:
        return {}, None

    base_level = min(block.level for _, block in headings)
    counters = [0, 0, 0]
    numbers: dict[int, str] = {}
    for index, block in headings:
        depth = block.level - base_level
        if depth < 0 or depth >= len(counters):
            continue
        for parent_depth in range(depth):
            if counters[parent_depth] == 0:
                counters[parent_depth] = 1
        counters[depth] += 1
        for deeper in range(depth + 1, len(counters)):
            counters[deeper] = 0
        numbers[index] = ".".join(
            str(counters[level]) for level in range(depth + 1)
        )
    return numbers, base_level


def has_manual_heading_numbers(
    blocks: list[MarkdownBlock],
    heading_numbers: dict[int, str],
) -> bool:
    top_level_indexes = [
        index
        for index in heading_numbers
        if "." not in heading_numbers[index]
    ]
    if not top_level_indexes:
        return False
    return all(
        re.match(
            rf"^{re.escape(heading_numbers[index])}[.)]\s+",
            blocks[index].text,
        )
        for index in top_level_indexes
    )


def numbered_heading_text(number: str, text: str) -> str:
    suffix = "." if "." not in number else ""
    return f"{number}{suffix} {strip_manual_heading_prefix(text)}"


def strip_manual_heading_prefix(text: str) -> str:
    return MANUAL_HEADING_PREFIX_PATTERN.sub("", text, count=1)


def add_heading_numbering_definition(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_num_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_num_id))
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "multilevel")
    abstract.append(multi_level_type)

    for level in range(3):
        level_element = OxmlElement("w:lvl")
        level_element.set(qn("w:ilvl"), str(level))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level_element.append(start)

        number_format = OxmlElement("w:numFmt")
        number_format.set(qn("w:val"), "decimal")
        level_element.append(number_format)

        level_text = OxmlElement("w:lvlText")
        pattern = ".".join(f"%{index + 1}" for index in range(level + 1))
        if level == 0:
            pattern += "."
        level_text.set(qn("w:val"), pattern)
        level_element.append(level_text)

        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        level_element.append(suffix)

        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        level_element.append(justification)

        paragraph_properties = OxmlElement("w:pPr")
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "0")
        indent.set(qn("w:hanging"), "0")
        paragraph_properties.append(indent)
        level_element.append(paragraph_properties)
        abstract.append(level_element)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_num_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def apply_heading_number(
    paragraph,
    *,
    num_id: int,
    level: int,
) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    number_properties = paragraph_properties.find(qn("w:numPr"))
    if number_properties is None:
        number_properties = OxmlElement("w:numPr")
        paragraph_properties.append(number_properties)

    indentation_level = OxmlElement("w:ilvl")
    indentation_level.set(qn("w:val"), str(level))
    number_properties.append(indentation_level)

    number_id = OxmlElement("w:numId")
    number_id.set(qn("w:val"), str(num_id))
    number_properties.append(number_id)


def restart_numbering_id(doc: Document, style_name: str) -> int:
    numbering = doc.part.numbering_part.element
    style_num_id = int(
        doc.styles[style_name]
        ._element.pPr.numPr.numId.get(qn("w:val"))
    )
    base_number = next(
        element
        for element in numbering.findall(qn("w:num"))
        if int(element.get(qn("w:numId"))) == style_num_id
    )
    abstract_num_id = base_number.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_num_id))
    number.append(abstract_reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    number.append(level_override)
    numbering.append(number)
    return num_id


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.69)
    section.bottom_margin = Inches(0.69)
    section.left_margin = Inches(0.69)
    section.right_margin = Inches(0.69)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for list_style_name in ("List Bullet", "List Number"):
        list_style = doc.styles[list_style_name]
        list_style.font.name = BODY_FONT
        list_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        list_style.font.size = Pt(9.5)
        list_style.paragraph_format.left_indent = Inches(0.32)
        list_style.paragraph_format.first_line_indent = Inches(-0.18)
        list_style.paragraph_format.space_after = Pt(4)
        list_style.paragraph_format.line_spacing = 1.08

    quote = doc.styles["Quote"]
    quote.font.name = BODY_FONT
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    quote.font.size = Pt(9.5)
    quote.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.10)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(5)
    quote.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE_DARK, 20, 10),
        ("Heading 2", 13, BLUE, 15, 7),
        ("Heading 3", 11, BLUE_DARK, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    clear_paragraphs(header.paragraphs)

    footer = section.footer
    footer.is_linked_to_previous = False
    clear_paragraphs(footer.paragraphs)
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.add_run("- ")
    add_field(footer_paragraph, "PAGE")
    footer_paragraph.add_run(" -")
    for run in footer_paragraph.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("999999")


def clear_paragraphs(paragraphs) -> None:
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)


def add_cover(
    doc: Document,
    document_title: str,
    saved_date: str,
    document_type: str = "",
    *,
    document_language: str = "ko",
) -> None:
    for _ in range(5):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(wrap_cover_title(document_title))
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE_DARK)

    if document_type:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(document_type)
        subtitle_run.font.name = BODY_FONT
        subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        subtitle_run.font.size = Pt(17)
        subtitle_run.font.bold = True
        subtitle_run.font.color.rgb = RGBColor.from_string(BLUE)

    line = doc.add_paragraph()
    add_bottom_border(line, BLUE, "8")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recorded_label = DOCUMENT_UI_LABELS[document_language]["recorded"]
    meta_run = meta.add_run(f"{recorded_label}: {saved_date}")
    meta_run.font.name = BODY_FONT
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor.from_string("666666")

    doc.add_page_break()


def wrap_cover_title(document_title: str, *, max_line_width: int = 34) -> str:
    """Wrap a long cover title only at a semantic word boundary."""
    title = " ".join(document_title.split())
    words = title.split(" ")
    if len(words) < 2 or display_width(title) <= max_line_width:
        return title

    candidates: list[tuple[int, int, int]] = []
    for index in range(1, len(words)):
        left_width = display_width(" ".join(words[:index]))
        right_width = display_width(" ".join(words[index:]))
        candidates.append(
            (
                abs(left_width - right_width),
                max(left_width, right_width),
                index,
            )
        )
    _, _, split_index = min(candidates)
    return (
        " ".join(words[:split_index])
        + "\n"
        + " ".join(words[split_index:])
    )


def display_width(text: str) -> int:
    return sum(
        2 if unicodedata.east_asian_width(character) in {"W", "F"} else 1
        for character in text
    )


def add_static_toc(
    doc: Document,
    blocks: list[MarkdownBlock],
    heading_anchors: dict[int, str],
    *,
    heading_numbers: dict[int, str] | None = None,
    title_index: int | None = None,
    document_language: str = "ko",
) -> None:
    heading_numbers = heading_numbers or {}
    contents_label = DOCUMENT_UI_LABELS[document_language]["contents"]
    heading = doc.add_paragraph(contents_label, style="Heading 1")
    add_bottom_border(heading, BLUE, "6")
    candidate_indexes = [
        index
        for index, block in enumerate(blocks)
        if block.kind == "heading"
        and block.level <= 3
        and index != title_index
    ]
    visible_indexes = set(candidate_indexes)
    if len(candidate_indexes) > TOC_DETAIL_ENTRY_LIMIT:
        top_level = min(blocks[index].level for index in candidate_indexes)
        visible_indexes = {
            index
            for index in candidate_indexes
            if blocks[index].level == top_level
        }
    for index, block in enumerate(blocks):
        if index not in visible_indexes:
            continue
        text = (
            numbered_heading_text(heading_numbers[index], block.text)
            if index in heading_numbers
            else block.text
        )
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.22 * max(block.level - 1, 0))
        add_internal_hyperlink(
            paragraph,
            text,
            heading_anchors[index],
            BLUE_DARK if block.level == 2 else "333333",
        )
    doc.add_page_break()


def add_body(
    doc: Document,
    blocks: list[MarkdownBlock],
    heading_anchors: dict[int, str],
    *,
    heading_num_id: int | None = None,
    numbering_base_level: int | None = None,
    numbered_heading_indexes: set[int] | None = None,
    title_index: int | None = None,
    source_dir: Path,
) -> None:
    numbered_heading_indexes = numbered_heading_indexes or set()
    for index, block in enumerate(blocks):
        if block.kind == "heading":
            if index == title_index:
                continue
            style = "Heading 1" if block.level <= 2 else "Heading 2" if block.level == 3 else "Heading 3"
            heading_text = (
                strip_manual_heading_prefix(block.text)
                if index in numbered_heading_indexes
                else block.text
            )
            paragraph = doc.add_paragraph(heading_text, style=style)
            if (
                heading_num_id is not None
                and numbering_base_level is not None
                and index in numbered_heading_indexes
            ):
                apply_heading_number(
                    paragraph,
                    num_id=heading_num_id,
                    level=block.level - numbering_base_level,
                )
            anchor = heading_anchors.get(index)
            if anchor is not None:
                add_bookmark(paragraph, anchor, index + 1)
            if style == "Heading 1":
                add_bottom_border(paragraph, BLUE, "6")
        elif block.kind == "paragraph":
            add_paragraph(doc, block.text)
        elif block.kind == "bullets":
            for item in block.items or []:
                add_bullet(doc, item)
        elif block.kind == "numbered":
            list_num_id = restart_numbering_id(doc, "List Number")
            for item in block.items or []:
                add_numbered_item(doc, item, num_id=list_num_id)
        elif block.kind == "blockquote":
            add_blockquote(doc, block.text)
        elif block.kind == "image":
            add_image(doc, block, source_dir=source_dir)
        elif block.kind == "table" and block.table is not None:
            add_table(doc, block.table)


def add_image(doc: Document, block: MarkdownBlock, *, source_dir: Path) -> None:
    raw_path = Path(block.path)
    image_path = raw_path if raw_path.is_absolute() else source_dir / raw_path
    resolved_source = source_dir.resolve()
    resolved_image = image_path.resolve()
    try:
        resolved_image.relative_to(resolved_source)
    except ValueError as exc:
        raise ValueError(f"Markdown image must stay inside {resolved_source}: {block.path}") from exc
    if not resolved_image.is_file():
        raise FileNotFoundError(f"Markdown image does not exist: {resolved_image}")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    inline_shape = paragraph.add_run().add_picture(
        image_descriptor_for_docx(resolved_image),
        width=Inches(6.1),
    )
    alt_text = block.text or resolved_image.name
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", alt_text)


def image_descriptor_for_docx(image_path: Path) -> str | io.BytesIO:
    """Add a JFIF marker in memory for valid FFmpeg JPEGs python-docx rejects."""
    data = image_path.read_bytes()
    if data.startswith(b"\xff\xd8") and data[6:10] not in {b"JFIF", b"Exif"}:
        return io.BytesIO(data[:2] + JFIF_APP0_SEGMENT + data[2:])
    return str(image_path)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    add_inline_markdown(paragraph, text, size=9.5)


def add_bullet(doc: Document, text: str) -> None:
    checkbox = CHECKBOX_ITEM_PATTERN.match(text)
    if checkbox:
        add_checklist_item(
            doc,
            checkbox.group("text"),
            checked=checkbox.group("mark").lower() == "x",
        )
        return
    paragraph = doc.add_paragraph(style="List Bullet")
    add_inline_markdown(paragraph, text, size=9.5)


def add_checklist_item(doc: Document, text: str, *, checked: bool) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.32)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.08
    add_formatted_run(
        paragraph,
        "☑ " if checked else "☐ ",
        size=9.5,
        bold=False,
        color="000000",
    )
    add_inline_markdown(paragraph, text, size=9.5)


def add_numbered_item(doc: Document, text: str, *, num_id: int) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    apply_heading_number(paragraph, num_id=num_id, level=0)
    add_inline_markdown(paragraph, text, size=9.5)


def add_blockquote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Quote")
    add_inline_markdown(paragraph, text, size=9.5)


def add_inline_markdown(
    paragraph,
    text: str,
    *,
    size: float,
    bold: bool = False,
    color: str = "000000",
) -> None:
    text = DANGLING_EVIDENCE_PATTERN.sub("", text)
    text = HTML_BREAK_PATTERN.sub("\n", text)
    position = 0
    for match in INLINE_MARKDOWN_PATTERN.finditer(text):
        add_formatted_run(
            paragraph,
            text[position : match.start()],
            size=size,
            bold=bold,
            color=color,
        )
        if match.group("code") is not None:
            add_formatted_run(
                paragraph,
                match.group("code"),
                size=size,
                bold=bold,
                color=color,
                font=CODE_FONT,
            )
        elif match.group("link_text") is not None:
            add_external_hyperlink(
                paragraph,
                match.group("link_text"),
                match.group("link_url"),
                size=size,
            )
        elif match.group("strong") is not None:
            add_formatted_run(
                paragraph,
                match.group("strong"),
                size=size,
                bold=True,
                color=color,
            )
        else:
            add_formatted_run(
                paragraph,
                match.group("emphasis"),
                size=size,
                bold=bold,
                color=color,
                italic=True,
            )
        position = match.end()
    add_formatted_run(
        paragraph,
        text[position:],
        size=size,
        bold=bold,
        color=color,
    )


def add_formatted_run(
    paragraph,
    text: str,
    *,
    size: float,
    bold: bool,
    color: str,
    font: str = BODY_FONT,
    italic: bool = False,
) -> None:
    if not text:
        return
    segments = text.split("\n")
    for index, segment in enumerate(segments):
        if segment:
            run = paragraph.add_run(segment)
            run.font.name = font
            run._element.rPr.rFonts.set(qn("w:ascii"), font)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.color.rgb = RGBColor.from_string(color)
        if index < len(segments) - 1:
            paragraph.add_run().add_break()


def add_external_hyperlink(paragraph, text: str, url: str, *, size: float) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    run_properties.append(fonts)

    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(size_element)

    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), BLUE)
    run_properties.append(color_element)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    text_element = OxmlElement("w:t")
    text_element.text = text

    run.append(run_properties)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(paragraph, text: str, anchor: str, color: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    run_properties.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    run_properties.append(size)

    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    run_properties.append(color_element)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    text_element = OxmlElement("w:t")
    text_element.text = text

    run.append(run_properties)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), name)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))

    insert_position = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_position, bookmark_start)
    paragraph._p.append(bookmark_end)


def add_table(doc: Document, table: MarkdownTable) -> None:
    width = max(len(table.headers), 1)
    word_table = doc.add_table(rows=len(table.rows) + 1, cols=width)
    widths = column_widths(width)
    set_table_width(word_table, 9360, widths)
    set_table_borders(word_table)
    set_table_cell_margins(word_table, 80, 120)
    for idx, cell in enumerate(word_table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_fill(cell, BLUE_DARK)
        set_cell_text(
            cell,
            table.headers[idx],
            bold=True,
            color=WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(word_table.rows[0])
    set_row_cant_split(word_table.rows[0])

    for row_index, row_values in enumerate(table.rows, start=1):
        cells = word_table.rows[row_index].cells
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            if idx == 0:
                set_cell_fill(cell, BLUE_LIGHT)
            elif row_index % 2 == 0:
                set_cell_fill(cell, GRAY_LIGHT)
            set_cell_text(cell, row_values[idx] if idx < len(row_values) else "")
        set_row_cant_split(word_table.rows[row_index])
    doc.add_paragraph("")


def column_widths(count: int) -> list[int]:
    if count == 1:
        return [9360]
    if count == 2:
        return [3000, 6360]
    if count == 3:
        return [2600, 3380, 3380]
    if count == 4:
        return [2300, 3200, 1800, 2060]
    if count == 5:
        # Five-column field guides commonly put a multi-word stage or category
        # in the first column and the longest narrative in the second.
        return [1400, 2960, 1800, 1400, 1800]
    total = 9360
    first = 800
    rest = int((total - first) / (count - 1))
    return [first] + [rest] * (count - 1)


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    color: str = "000000",
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    clear_paragraphs([paragraph])
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.10
    add_inline_markdown(
        paragraph,
        text,
        size=8.5,
        bold=bold,
        color=color,
    )


def set_table_width(table, width: int, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for col_width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(col_width))
        grid.append(grid_col)

    for idx, col_width in enumerate(widths):
        table.columns[idx].width = Twips(col_width)


def set_cell_width(cell, width: int) -> None:
    cell.width = Twips(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width))
    set_cell_margins(cell, 80, 120)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_cell_margins(table, vertical: int, horizontal: int) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, margin in (
        ("top", vertical),
        ("bottom", vertical),
        ("start", horizontal),
        ("end", horizontal),
    ):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(margin))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_margins(cell, vertical: int, horizontal: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, margin in (
        ("top", vertical),
        ("bottom", vertical),
        ("start", horizontal),
        ("end", horizontal),
    ):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(margin))
        element.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color: str, size: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    border = p_pr.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        p_pr.append(border)
    bottom = border.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        border.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def main(argv: Sequence[str] | None = None) -> None:
    parser = argparse.ArgumentParser(
        description="Fill the retained minutes Word template from Markdown."
    )
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--title")
    parser.add_argument("--document-type")
    parser.add_argument("--saved-date", default="")
    parser.add_argument("--template", type=Path)
    args = parser.parse_args(list(argv) if argv is not None else None)
    result = generate_docx_report(
        args.markdown,
        args.output,
        document_title=args.title,
        document_type=args.document_type,
        saved_date=args.saved_date,
        template_path=args.template,
    )
    print(result)


if __name__ == "__main__":
    main()
