from __future__ import annotations

import argparse
import hashlib
import json
import re
import struct
import sys
from pathlib import Path
from typing import Any, Sequence
from zipfile import BadZipFile, ZipFile

if __package__ in {None, ""}:
    sys.path.append(str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.oxml.ns import qn

from scripts.docx_report import (
    DOCUMENTS_PRESET,
    HEADER_PATTERN,
    INTERNAL_CITATION_PATTERN,
    validate_supported_markdown,
)
from scripts.utils import now_local, read_json, write_json


SCHEMA_VERSION = 1
SELECTED_REVIEW_PRESET = DOCUMENTS_PRESET
EXPECTED_STYLES = {
    "Normal": {
        "font": "Calibri",
        "size_pt": 11.0,
        "before_pt": 0.0,
        "after_pt": 6.0,
        "line_spacing": 1.10,
    },
    "Heading 1": {
        "font": "Calibri",
        "size_pt": 16.0,
        "before_pt": 16.0,
        "after_pt": 8.0,
        "line_spacing": None,
    },
    "Heading 2": {
        "font": "Calibri",
        "size_pt": 13.0,
        "before_pt": 12.0,
        "after_pt": 6.0,
        "line_spacing": None,
    },
    "Heading 3": {
        "font": "Calibri",
        "size_pt": 12.0,
        "before_pt": 8.0,
        "after_pt": 4.0,
        "line_spacing": None,
    },
    "List Bullet": {
        "font": "Calibri",
        "size_pt": 11.0,
        "before_pt": None,
        "after_pt": 8.0,
        "line_spacing": 1.167,
    },
    "List Number": {
        "font": "Calibri",
        "size_pt": 11.0,
        "before_pt": None,
        "after_pt": 8.0,
        "line_spacing": 1.167,
    },
}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def file_record(path: Path) -> dict[str, Any]:
    return {
        "path": str(path),
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
    }


def audit_docx(docx_path: Path) -> dict[str, Any]:
    issues: list[str] = []
    try:
        with ZipFile(docx_path) as archive:
            corrupt_member = archive.testzip()
    except (OSError, BadZipFile) as exc:
        return {
            "passed": False,
            "issues": [f"invalid DOCX package: {exc}"],
            "preset": SELECTED_REVIEW_PRESET,
        }
    if corrupt_member is not None:
        issues.append(f"corrupt DOCX member: {corrupt_member}")

    document = Document(docx_path)
    preset_audit = _audit_preset(document, issues)
    structure_audit = _audit_structure(document, issues)
    marker_audit = _audit_literal_markers(document, issues)
    return {
        "passed": not issues,
        "issues": issues,
        "preset": SELECTED_REVIEW_PRESET,
        "header_pattern": HEADER_PATTERN,
        "preset_audit": preset_audit,
        "structure_audit": structure_audit,
        "marker_audit": marker_audit,
    }


def _audit_preset(document: Any, issues: list[str]) -> dict[str, Any]:
    section = document.sections[0]
    page = {
        "width_in": round(section.page_width.inches, 3),
        "height_in": round(section.page_height.inches, 3),
        "top_margin_in": round(section.top_margin.inches, 3),
        "right_margin_in": round(section.right_margin.inches, 3),
        "bottom_margin_in": round(section.bottom_margin.inches, 3),
        "left_margin_in": round(section.left_margin.inches, 3),
        "header_distance_in": round(section.header_distance.inches, 3),
        "footer_distance_in": round(section.footer_distance.inches, 3),
    }
    expected_page = {
        "width_in": 8.5,
        "height_in": 11.0,
        "top_margin_in": 1.0,
        "right_margin_in": 1.0,
        "bottom_margin_in": 1.0,
        "left_margin_in": 1.0,
        "header_distance_in": 0.492,
        "footer_distance_in": 0.492,
    }
    for key, expected in expected_page.items():
        if abs(page[key] - expected) > 0.005:
            issues.append(f"preset page token mismatch: {key}={page[key]} expected {expected}")

    styles: dict[str, dict[str, Any]] = {}
    for style_name, expected in EXPECTED_STYLES.items():
        style = document.styles[style_name]
        actual = {
            "font": style.font.name,
            "size_pt": _points(style.font.size),
            "before_pt": _points(style.paragraph_format.space_before),
            "after_pt": _points(style.paragraph_format.space_after),
            "line_spacing": _line_spacing(style.paragraph_format.line_spacing),
        }
        styles[style_name] = actual
        for key, expected_value in expected.items():
            if expected_value is None:
                continue
            actual_value = actual[key]
            if isinstance(expected_value, float) and isinstance(actual_value, (int, float)):
                if abs(float(actual_value) - expected_value) <= 0.01:
                    continue
            elif actual_value == expected_value:
                continue
            issues.append(
                f"preset style token mismatch: {style_name}.{key}="
                f"{actual_value!r} expected {expected_value!r}"
            )

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        if style._element.pPr is None or style._element.pPr.numPr is None:
            issues.append(f"{style_name} does not use a real numbering definition")
        left_indent = style.paragraph_format.left_indent
        first_line = style.paragraph_format.first_line_indent
        if left_indent is None or abs(left_indent.inches - 0.5) > 0.005:
            issues.append(f"{style_name} text indent does not match 0.5in")
        if first_line is None or abs(first_line.inches + 0.25) > 0.005:
            issues.append(f"{style_name} hanging indent does not match 0.25in")
    return {"page": page, "styles": styles}


def _audit_structure(document: Any, issues: list[str]) -> dict[str, Any]:
    root = document._element
    links = root.findall(".//" + qn("w:hyperlink"))
    internal_anchors = {
        link.get(qn("w:anchor"))
        for link in links
        if link.get(qn("w:anchor"))
    }
    bookmark_names = {
        bookmark.get(qn("w:name"))
        for bookmark in root.findall(".//" + qn("w:bookmarkStart"))
    }
    missing_bookmarks = sorted(internal_anchors - bookmark_names)
    if missing_bookmarks:
        issues.append(f"TOC anchors lack matching bookmarks: {missing_bookmarks}")

    table_results = []
    for index, table in enumerate(document.tables):
        table_issues = _audit_table_geometry(table)
        if table_issues:
            issues.extend(f"table[{index}] {issue}" for issue in table_issues)
        table_results.append(
            {
                "index": index,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "issues": table_issues,
            }
        )
    return {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "internal_link_count": len(internal_anchors),
        "bookmark_count": len(bookmark_names),
        "missing_bookmarks": missing_bookmarks,
        "tables": table_results,
    }


def _audit_table_geometry(table: Any) -> list[str]:
    issues: list[str] = []
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_w is None or tbl_w.get(qn("w:type")) != "dxa" or tbl_w.get(qn("w:w")) != "9360":
        issues.append("tblW must be 9360 DXA")
    if tbl_ind is None or tbl_ind.get(qn("w:type")) != "dxa" or tbl_ind.get(qn("w:w")) != "120":
        issues.append("tblInd must be 120 DXA")
    grid_widths = [
        int(column.get(qn("w:w"), "0"))
        for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
    ]
    if not grid_widths or sum(grid_widths) != 9360:
        issues.append("tblGrid widths must sum to 9360 DXA")
    for row_index, row in enumerate(table.rows):
        cell_widths = []
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            cell_widths.append(int(tc_w.get(qn("w:w"), "0")) if tc_w is not None else 0)
        if cell_widths != grid_widths:
            issues.append(
                f"row {row_index} tcW values {cell_widths} do not match tblGrid {grid_widths}"
            )
        tr_height = row._tr.get_or_add_trPr().find(qn("w:trHeight"))
        if tr_height is not None and tr_height.get(qn("w:hRule")) == "exact":
            issues.append(f"row {row_index} uses an exact fixed height")
    if table.rows:
        header = table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader"))
        if header is None:
            issues.append("first row is not marked as a repeating header")
    return issues


def _audit_literal_markers(document: Any, issues: list[str]) -> dict[str, Any]:
    marker_hits: list[str] = []
    fake_list_hits: list[str] = []
    body_paragraphs = list(document.paragraphs)
    all_paragraphs = list(body_paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for paragraph in all_paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if (
            "**" in text
            or text.startswith("> ")
            or re.match(r"^\[[ xX]\]\s+", text)
            or INTERNAL_CITATION_PATTERN.search(text)
        ):
            marker_hits.append(text[:160])
    for index, paragraph in enumerate(body_paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        is_list_style = paragraph.style.name in {"List Bullet", "List Number"}
        is_toc_link = paragraph._p.find(qn("w:hyperlink")) is not None
        if is_list_style or is_toc_link:
            continue
        if re.match(r"^[-*•]\s+", text):
            fake_list_hits.append(text[:160])
            continue
        if not re.match(r"^\d+[.)]\s+", text):
            continue
        neighbors = body_paragraphs[max(0, index - 1) : index] + body_paragraphs[
            index + 1 : index + 2
        ]
        if any(
            neighbor.style.name not in {"List Bullet", "List Number"}
            and re.match(r"^\d+[.)]\s+", neighbor.text.strip())
            for neighbor in neighbors
        ):
            fake_list_hits.append(text[:160])
    if marker_hits:
        issues.append(f"literal Markdown or internal citation markers remain: {marker_hits}")
    if fake_list_hits:
        issues.append(f"fake list paragraphs remain: {fake_list_hits}")
    return {
        "literal_marker_hits": marker_hits,
        "fake_list_hits": fake_list_hits,
    }


def render_record(render_dir: Path, *, visual_status: str) -> dict[str, Any]:
    page_paths = sorted(render_dir.glob("page-*.png"), key=_page_number)
    pages = []
    for path in page_paths:
        width, height = png_dimensions(path)
        pages.append(
            {
                "page": _page_number(path),
                "path": str(path),
                "bytes": path.stat().st_size,
                "width": width,
                "height": height,
                "sha256": sha256_file(path),
            }
        )
    manifest_payload = json.dumps(
        [
            {"page": page["page"], "sha256": page["sha256"]}
            for page in pages
        ],
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return {
        "visual_inspection": visual_status,
        "page_count": len(pages),
        "manifest_sha256": hashlib.sha256(manifest_payload).hexdigest(),
        "pages": pages,
    }


def _visual_review_issues(
    visual_review: Any,
    *,
    final_docx_sha256: str,
    rendered: dict[str, Any],
) -> list[str]:
    issues: list[str] = []
    if not isinstance(visual_review, dict):
        return ["visual review must be an object"]
    if visual_review.get("schema_version") != 1:
        issues.append("visual review schema_version must be 1")
    if visual_review.get("status") != "passed":
        issues.append("visual review status must be passed")
    expected_pages = list(range(1, int(rendered.get("page_count", 0)) + 1))
    if visual_review.get("inspected_pages") != expected_pages:
        issues.append("visual review must inspect every latest rendered page")
    blockers = visual_review.get("blocking_defects")
    if not isinstance(blockers, list) or blockers:
        issues.append("visual review blocking_defects must be an empty list")
    warnings = visual_review.get("warnings")
    if not isinstance(warnings, list):
        issues.append("visual review warnings must be a list")
    bindings = visual_review.get("bindings")
    if not isinstance(bindings, dict):
        issues.append("visual review bindings must be an object")
    else:
        if bindings.get("final_docx_sha256") != final_docx_sha256:
            issues.append("visual review final DOCX hash does not match")
        if bindings.get("render_manifest_sha256") != rendered.get("manifest_sha256"):
            issues.append("visual review render manifest hash does not match")
    return issues


def _current_render_issues(rendered: Any) -> list[str]:
    if not isinstance(rendered, dict):
        return ["render record must be an object"]
    pages = rendered.get("pages")
    if not isinstance(pages, list):
        return ["render pages must be a list"]
    issues: list[str] = []
    current_records: list[dict[str, Any]] = []
    for index, page in enumerate(pages):
        if not isinstance(page, dict):
            issues.append(f"render pages[{index}] must be an object")
            continue
        path = Path(str(page.get("path", ""))).expanduser()
        if not path.is_file():
            issues.append(f"rendered page is missing: {path}")
            continue
        try:
            width, height = png_dimensions(path)
        except ValueError as exc:
            issues.append(str(exc))
            continue
        current = {
            "page": _page_number(path),
            "path": str(path),
            "bytes": path.stat().st_size,
            "width": width,
            "height": height,
            "sha256": sha256_file(path),
        }
        current_records.append(current)
        if current != page:
            issues.append(f"rendered page changed after review: {path.name}")
    current_manifest = hashlib.sha256(
        json.dumps(
            [
                {"page": page["page"], "sha256": page["sha256"]}
                for page in current_records
            ],
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()
    if rendered.get("page_count") != len(current_records):
        issues.append("render page count no longer matches current page files")
    if (
        rendered.get("manifest_sha256") is not None
        and rendered.get("manifest_sha256") != current_manifest
    ):
        issues.append("render manifest no longer matches current page files")
    return issues


def create_docx_qa(
    markdown_path: Path,
    draft_docx_path: Path,
    final_docx_path: Path,
    *,
    render_dir: Path,
    visual_status: str,
    output_path: Path,
    renderer: str = "documents/render_docx.py",
    visual_review: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if visual_status not in {"passed", "failed", "not_run"}:
        raise ValueError("visual_status must be passed, failed, or not_run")
    for path in (markdown_path, draft_docx_path, final_docx_path):
        if not path.is_file():
            raise FileNotFoundError(path)
    markdown = markdown_path.read_text(encoding="utf-8")
    markdown_issues: list[str] = []
    try:
        validate_supported_markdown(markdown)
    except ValueError as exc:
        markdown_issues.append(str(exc))
    structural = audit_docx(final_docx_path)
    rendered = render_record(render_dir, visual_status=visual_status)
    issues = list(markdown_issues) + list(structural["issues"])
    if visual_status == "passed" and rendered["page_count"] == 0:
        issues.append("visual inspection cannot pass without rendered page PNGs")
    if visual_status == "failed":
        issues.append("visual inspection failed")
    final_docx_record = file_record(final_docx_path)
    if visual_status == "passed" and visual_review is not None:
        issues.extend(
            _visual_review_issues(
                visual_review,
                final_docx_sha256=final_docx_record["sha256"],
                rendered=rendered,
            )
        )
    status = (
        "passed"
        if not issues and visual_status == "passed"
        else "structural_only"
        if not issues and visual_status == "not_run"
        else "failed"
    )
    result = {
        "schema_version": SCHEMA_VERSION,
        "status": status,
        "checked_at": now_local().isoformat(),
        "documents_preset": SELECTED_REVIEW_PRESET,
        "header_pattern": HEADER_PATTERN,
        "renderer": renderer,
        "source_markdown": file_record(markdown_path),
        "draft_docx": file_record(draft_docx_path),
        "final_docx": final_docx_record,
        "markdown_preflight": {
            "passed": not markdown_issues,
            "issues": markdown_issues,
        },
        "structural": structural,
        "render": rendered,
        "visual_review": visual_review,
        "issues": issues,
    }
    write_json(output_path, result)
    return result


def validate_docx_qa(
    markdown_path: Path,
    final_docx_path: Path,
    qa_path: Path,
    *,
    require_visual: bool,
    require_visual_review: bool = False,
) -> dict[str, Any]:
    qa = read_json(qa_path)
    issues: list[str] = []
    if qa.get("schema_version") != SCHEMA_VERSION:
        issues.append("docx_qa.json schema_version must be 1")
    expected_status = "passed" if require_visual else None
    if expected_status is not None and qa.get("status") != expected_status:
        issues.append("docx_qa.json status must be passed")
    if not require_visual and qa.get("status") not in {"passed", "structural_only"}:
        issues.append("docx_qa.json status must be passed or structural_only")
    if qa.get("documents_preset") != SELECTED_REVIEW_PRESET:
        issues.append(f"docx_qa.json must use {SELECTED_REVIEW_PRESET}")
    source_record = qa.get("source_markdown", {})
    final_record = qa.get("final_docx", {})
    if not isinstance(source_record, dict) or source_record.get("sha256") != sha256_file(markdown_path):
        issues.append("docx_qa.json source Markdown hash does not match")
    if not isinstance(final_record, dict) or final_record.get("sha256") != sha256_file(final_docx_path):
        issues.append("docx_qa.json final DOCX hash does not match")
    render = qa.get("render", {})
    structural_record = qa.get("structural", {})
    if not isinstance(structural_record, dict) or structural_record.get("passed") is not True:
        issues.append("docx_qa.json structural audit did not pass")
    if require_visual and (
        not isinstance(render, dict)
        or render.get("visual_inspection") != "passed"
        or int(render.get("page_count", 0)) <= 0
    ):
        issues.append("docx_qa.json requires passed full-page render inspection")
    if require_visual:
        issues.extend(
            "docx_qa.json " + issue for issue in _current_render_issues(render)
        )
    visual_review = qa.get("visual_review")
    if require_visual_review and visual_review is None:
        issues.append("docx_qa.json requires a hash-bound all-page visual review")
    if require_visual and visual_review is not None and isinstance(render, dict):
        issues.extend(
            "docx_qa.json " + issue
            for issue in _visual_review_issues(
                visual_review,
                final_docx_sha256=sha256_file(final_docx_path),
                rendered=render,
            )
        )
    current_structural = audit_docx(final_docx_path)
    if not current_structural["passed"]:
        issues.extend(current_structural["issues"])
    if issues:
        raise ValueError("DOCX QA failed: " + "; ".join(issues))
    return qa


def png_dimensions(path: Path) -> tuple[int, int]:
    with path.open("rb") as handle:
        header = handle.read(24)
    if len(header) < 24 or header[:8] != b"\x89PNG\r\n\x1a\n" or header[12:16] != b"IHDR":
        raise ValueError(f"invalid PNG page render: {path}")
    return struct.unpack(">II", header[16:24])


def _page_number(path: Path) -> int:
    match = re.search(r"page-(\d+)\.png$", path.name)
    return int(match.group(1)) if match else 0


def _points(value: Any) -> float | None:
    return round(value.pt, 3) if value is not None else None


def _line_spacing(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return round(float(value), 3)
    return _points(value)


def main(argv: Sequence[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description="Create a hash-bound DOCX QA artifact.")
    parser.add_argument("--markdown", type=Path, required=True)
    parser.add_argument("--draft", type=Path, required=True)
    parser.add_argument("--final", type=Path, required=True)
    parser.add_argument("--render-dir", type=Path, required=True)
    parser.add_argument(
        "--visual-status",
        choices=("passed", "failed", "not_run"),
        required=True,
    )
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--renderer", default="documents/render_docx.py")
    parser.add_argument(
        "--visual-review",
        type=Path,
        help="Optional hash-bound visual_review.json produced after all-page inspection",
    )
    args = parser.parse_args(list(argv) if argv is not None else None)
    result = create_docx_qa(
        args.markdown,
        args.draft,
        args.final,
        render_dir=args.render_dir,
        visual_status=args.visual_status,
        output_path=args.output,
        renderer=args.renderer,
        visual_review=read_json(args.visual_review) if args.visual_review else None,
    )
    print(
        json.dumps(
            {
                "output": str(args.output),
                "status": result["status"],
                "preset": result["documents_preset"],
                "page_count": result["render"]["page_count"],
                "final_docx_sha256": result["final_docx"]["sha256"],
                "issues": result["issues"],
            },
            ensure_ascii=False,
        )
    )
    if result["status"] == "failed":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
