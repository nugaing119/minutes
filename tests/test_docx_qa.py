from __future__ import annotations

import shutil
import struct
import subprocess
import tempfile
import unittest
from pathlib import Path

from docx import Document

from scripts.docx_qa import (
    audit_docx,
    create_docx_qa,
    render_record,
    sha256_file,
    validate_docx_qa,
)
from scripts.docx_report import generate_docx_report


def write_png_header(path: Path, width: int = 1275, height: int = 1650) -> None:
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + struct.pack(">I", 13)
        + b"IHDR"
        + struct.pack(">II", width, height)
    )


class DocxQaTests(unittest.TestCase):
    def test_script_entrypoint_resolves_repo_imports_without_pythonpath(self) -> None:
        repo_root = Path(__file__).resolve().parents[1]
        result = subprocess.run(
            [str(repo_root / ".venv/bin/python"), "scripts/docx_qa.py", "--help"],
            cwd=repo_root,
            capture_output=True,
            text=True,
            check=False,
        )

        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertIn("Create a hash-bound DOCX QA artifact", result.stdout)

    def _artifacts(self, root: Path) -> tuple[Path, Path, Path, Path]:
        markdown = root / "minutes.md"
        draft = root / "minutes.draft.docx"
        final = root / "minutes.final.docx"
        render_dir = root / "docx_render"
        render_dir.mkdir()
        markdown.write_text(
            "# 기술 검토\n\n"
            "문서 유형: 기술 브리프\n\n"
            "## 조치\n\n"
            "**중요:** 근거를 보존한다.\n\n"
            "1. 검증한다.\n"
            "2. 아카이브한다.\n\n"
            "| 항목 | 상태 |\n"
            "|---|---|\n"
            "| 증거 | 통과 |\n",
            encoding="utf-8",
        )
        generate_docx_report(markdown, draft, saved_date="2026-07-15")
        shutil.copy2(draft, final)
        return markdown, draft, final, render_dir

    def test_hash_bound_full_render_qa_passes_and_revalidates(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown, draft, final, render_dir = self._artifacts(root)
            write_png_header(render_dir / "page-1.png")
            qa_path = root / "docx_qa.json"

            qa = create_docx_qa(
                markdown,
                draft,
                final,
                render_dir=render_dir,
                visual_status="passed",
                output_path=qa_path,
            )
            verified = validate_docx_qa(
                markdown,
                final,
                qa_path,
                require_visual=True,
            )

        self.assertEqual(qa["status"], "passed")
        self.assertEqual(qa["documents_preset"], "retained_word_template")
        self.assertTrue(qa["structural"]["passed"])
        self.assertEqual(qa["render"]["page_count"], 1)
        self.assertEqual(verified["final_docx"]["sha256"], qa["final_docx"]["sha256"])

    def test_final_docx_hash_change_invalidates_existing_qa(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown, draft, final, render_dir = self._artifacts(root)
            write_png_header(render_dir / "page-1.png")
            qa_path = root / "docx_qa.json"
            create_docx_qa(
                markdown,
                draft,
                final,
                render_dir=render_dir,
                visual_status="passed",
                output_path=qa_path,
            )
            document = Document(final)
            document.add_paragraph("changed after review")
            document.save(final)

            with self.assertRaisesRegex(ValueError, "final DOCX hash does not match"):
                validate_docx_qa(
                    markdown,
                    final,
                    qa_path,
                    require_visual=True,
                )

    def test_visual_pass_is_rejected_when_no_page_png_exists(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown, draft, final, render_dir = self._artifacts(root)

            qa = create_docx_qa(
                markdown,
                draft,
                final,
                render_dir=render_dir,
                visual_status="passed",
                output_path=root / "docx_qa.json",
            )

        self.assertEqual(qa["status"], "failed")
        self.assertIn("without rendered page PNGs", qa["issues"][0])

    def test_ooxml_marker_audit_detects_literal_markdown_leak(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            _markdown, _draft, final, _render_dir = self._artifacts(root)
            document = Document(final)
            document.add_paragraph("**literal leak**")
            document.save(final)

            audit = audit_docx(final)

        self.assertFalse(audit["passed"])
        self.assertTrue(audit["marker_audit"]["literal_marker_hits"])

    def test_ooxml_marker_audit_detects_literal_checklist_leak(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            _markdown, _draft, final, _render_dir = self._artifacts(root)
            document = Document(final)
            document.add_paragraph("[ ] unresolved action")
            document.save(final)

            audit = audit_docx(final)

        self.assertFalse(audit["passed"])
        self.assertTrue(audit["marker_audit"]["literal_marker_hits"])

    def test_table_stage_labels_are_not_treated_as_fake_body_lists(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            _markdown, _draft, final, _render_dir = self._artifacts(root)
            document = Document(final)
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Stage"
            table.cell(1, 0).text = "1. Introduce"
            document.save(final)

            audit = audit_docx(final)

        self.assertFalse(
            any("1. Introduce" in item for item in audit["marker_audit"]["fake_list_hits"])
        )

    def test_consecutive_fake_numbered_body_paragraphs_are_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            _markdown, _draft, final, _render_dir = self._artifacts(root)
            document = Document(final)
            document.add_paragraph("1. first fake item")
            document.add_paragraph("2. second fake item")
            document.save(final)

            audit = audit_docx(final)

        self.assertFalse(audit["passed"])
        self.assertEqual(len(audit["marker_audit"]["fake_list_hits"]), 2)

    def test_hash_bound_visual_review_requires_every_page(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown, draft, final, render_dir = self._artifacts(root)
            write_png_header(render_dir / "page-1.png")
            rendered = render_record(render_dir, visual_status="passed")
            review = {
                "schema_version": 1,
                "status": "passed",
                "inspected_pages": [1],
                "blocking_defects": [],
                "warnings": [],
                "bindings": {
                    "final_docx_sha256": sha256_file(final),
                    "render_manifest_sha256": rendered["manifest_sha256"],
                },
            }

            qa = create_docx_qa(
                markdown,
                draft,
                final,
                render_dir=render_dir,
                visual_status="passed",
                output_path=root / "docx_qa.json",
                visual_review=review,
            )

        self.assertEqual(qa["status"], "passed")
        self.assertEqual(qa["visual_review"]["inspected_pages"], [1])

    def test_new_frozen_workflow_cannot_skip_visual_review_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown, draft, final, render_dir = self._artifacts(root)
            write_png_header(render_dir / "page-1.png")
            qa_path = root / "docx_qa.json"
            create_docx_qa(
                markdown,
                draft,
                final,
                render_dir=render_dir,
                visual_status="passed",
                output_path=qa_path,
            )

            with self.assertRaisesRegex(ValueError, "all-page visual review"):
                validate_docx_qa(
                    markdown,
                    final,
                    qa_path,
                    require_visual=True,
                    require_visual_review=True,
                )

    def test_render_page_change_invalidates_existing_qa(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown, draft, final, render_dir = self._artifacts(root)
            page = render_dir / "page-1.png"
            write_png_header(page)
            qa_path = root / "docx_qa.json"
            create_docx_qa(
                markdown,
                draft,
                final,
                render_dir=render_dir,
                visual_status="passed",
                output_path=qa_path,
            )
            page.write_bytes(page.read_bytes() + b"changed")

            with self.assertRaisesRegex(ValueError, "rendered page changed"):
                validate_docx_qa(
                    markdown,
                    final,
                    qa_path,
                    require_visual=True,
                )


if __name__ == "__main__":
    unittest.main()
