from __future__ import annotations

import base64
import tempfile
import unittest
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

from scripts.docx_report import (
    DOCUMENTS_PRESET,
    column_widths,
    generate_docx_report,
    image_descriptor_for_docx,
    wrap_cover_title,
)


class DocxReportTests(unittest.TestCase):
    def test_ffmpeg_jpeg_without_jfif_marker_is_normalized_in_memory(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            image_path = Path(temp_dir) / "frame.jpg"
            original = b"\xff\xd8\xff\xfe\x00\x08Lavc\x00\xff\xd9"
            image_path.write_bytes(original)

            descriptor = image_descriptor_for_docx(image_path)

        self.assertNotEqual(descriptor, image_path)
        normalized = descriptor.getvalue()
        self.assertEqual(normalized[:2], b"\xff\xd8")
        self.assertEqual(normalized[6:10], b"JFIF")
        self.assertTrue(normalized.endswith(original[2:]))

    def test_front_matter_images_html_breaks_and_table_rows_render_semantically(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            snapshots = root / "snapshots"
            snapshots.mkdir()
            image_path = snapshots / "snapshot_0001_00-00-10.png"
            image_path.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZGU0AAAAASUVORK5CYII="
                )
            )
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "---\n"
                'title: "Customer journey"\n'
                'document_type: "Internal field guide"\n'
                'output_language: "English"\n'
                "---\n\n"
                "# Customer journey\n\n"
                "## Journey\n\n"
                "![Five-stage journey](snapshots/snapshot_0001_00-00-10.png)\n\n"
                "*Five stages. Evidence:.*\n\n"
                "| Stage | Evidence |\n"
                "|---|---|\n"
                "| Introduce | `STT:00:00:01-00:00:02`<br>`OCR:00:00:10` |\n",
                encoding="utf-8",
            )

            generate_docx_report(markdown, output, saved_date="2026-07-15")
            document = Document(output)

        paragraph_text = [paragraph.text for paragraph in document.paragraphs]
        all_text = "\n".join(
            paragraph_text
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )
        self.assertFalse(any(text in {"---", 'title: "Customer journey"'} for text in paragraph_text))
        self.assertNotIn("<br>", all_text)
        self.assertNotIn("Evidence:.", all_text)
        self.assertIn("STT:00:00:01-00:00:02\nOCR:00:00:10", all_text)
        self.assertEqual(len(document.inline_shapes), 1)
        self.assertEqual(
            document.inline_shapes[0]._inline.docPr.get("descr"),
            "Five-stage journey",
        )
        for row in document.tables[0].rows:
            self.assertIsNotNone(row._tr.get_or_add_trPr().find(qn("w:cantSplit")))

    def test_mixed_manual_subheadings_and_separate_lists_do_not_double_or_continue(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "# Guide\n\n"
                "Document type: Field guide\n\n"
                "## Workflow\n\n"
                "### 1. Prepare\n\n"
                "1. First preparation step.\n"
                "2. Second preparation step.\n\n"
                "### 2. Deliver\n\n"
                "1. First delivery step.\n"
                "2. Second delivery step.\n",
                encoding="utf-8",
            )

            generate_docx_report(markdown, output, saved_date="2026-07-15")
            document = Document(output)

            namespace = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }
            with ZipFile(output) as archive:
                root_xml = ET.fromstring(archive.read("word/document.xml"))
            toc_texts = [
                "".join(link.itertext())
                for link in root_xml.findall(".//w:hyperlink", namespace)
            ]

        self.assertEqual(toc_texts, ["1. Workflow", "1.1 Prepare", "1.2 Deliver"])
        heading_text = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.style.name in {"Heading 1", "Heading 2"}
            and paragraph.text != "Contents"
        ]
        self.assertEqual(heading_text, ["Workflow", "Prepare", "Deliver"])
        numbered = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name == "List Number"
        ]
        self.assertEqual(len(numbered), 4)
        num_ids = [
            paragraph._p.pPr.numPr.numId.get(qn("w:val"))
            for paragraph in numbered
        ]
        self.assertEqual(num_ids[0], num_ids[1])
        self.assertEqual(num_ids[2], num_ids[3])
        self.assertNotEqual(num_ids[0], num_ids[2])
        numbering = document.part.numbering_part.element
        for num_id in {num_ids[0], num_ids[2]}:
            number = next(
                element
                for element in numbering.findall(qn("w:num"))
                if element.get(qn("w:numId")) == num_id
            )
            self.assertEqual(
                number.find(qn("w:lvlOverride")).find(qn("w:startOverride")).get(qn("w:val")),
                "1",
            )

    def test_five_column_table_reserves_readable_first_column(self) -> None:
        widths = column_widths(5)

        self.assertEqual(sum(widths), 9360)
        self.assertGreaterEqual(widths[0], 900)

    def test_cover_and_contents_labels_follow_document_language(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            english_markdown = root / "english.md"
            english_output = root / "english.docx"
            english_markdown.write_text(
                "# Product briefing\n\n"
                "Document type: Internal briefing\n\n"
                "## Direction\n\n"
                "Body\n",
                encoding="utf-8",
            )
            generate_docx_report(
                english_markdown,
                english_output,
                saved_date="2026-04-13",
            )

            korean_markdown = root / "korean.md"
            korean_output = root / "korean.docx"
            korean_markdown.write_text(
                "# 제품 브리핑\n\n"
                "문서 유형: 내부 브리핑\n\n"
                "## 방향\n\n"
                "본문\n",
                encoding="utf-8",
            )
            generate_docx_report(
                korean_markdown,
                korean_output,
                saved_date="2026-04-13",
            )

            english_text = [
                paragraph.text for paragraph in Document(english_output).paragraphs
            ]
            korean_text = [
                paragraph.text for paragraph in Document(korean_output).paragraphs
            ]

        self.assertIn("Contents", english_text)
        self.assertIn("Recorded: 2026-04-13", english_text)
        self.assertNotIn("목차", english_text)
        self.assertFalse(any("작성일" in text for text in english_text))
        self.assertIn("목차", korean_text)
        self.assertIn("기록일: 2026-04-13", korean_text)
        self.assertFalse(any("작성일" in text for text in korean_text))

    def test_standard_business_brief_preset_is_encoded_in_styles_and_page(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "# 기술 검토\n\n문서 유형: 기술 브리프\n\n## 결론\n\n본문\n",
                encoding="utf-8",
            )

            generate_docx_report(markdown, output, saved_date="2026-07-15")
            document = Document(output)

        self.assertEqual(DOCUMENTS_PRESET, "standard_business_brief")
        section = document.sections[0]
        self.assertAlmostEqual(section.top_margin.inches, 1.0, places=3)
        self.assertAlmostEqual(section.right_margin.inches, 1.0, places=3)
        normal = document.styles["Normal"]
        self.assertEqual(normal.font.name, "Calibri")
        self.assertEqual(normal.font.size.pt, 11.0)
        self.assertEqual(normal.paragraph_format.space_after.pt, 6.0)
        self.assertEqual(document.styles["Heading 1"].font.size.pt, 16.0)
        self.assertEqual(document.styles["Heading 2"].font.size.pt, 13.0)
        self.assertEqual(document.styles["Heading 3"].font.size.pt, 12.0)

    def test_bold_blockquote_and_ordered_list_do_not_leak_markdown_markers(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "# 기술 검토\n\n"
                "문서 유형: 기술 브리프\n\n"
                "## 조치\n\n"
                "**중요:** 순서대로 수행한다.\n\n"
                "> 운영 중에는 기존 근거를 보존한다.\n\n"
                "1. 먼저 검증한다.\n"
                "2. 이후 배포한다.\n",
                encoding="utf-8",
            )

            generate_docx_report(markdown, output, saved_date="2026-07-15")
            document = Document(output)

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertNotIn("**", text)
        self.assertNotIn("> 운영", text)
        self.assertFalse(
            any(
                paragraph.text.startswith(("1. ", "2. "))
                for paragraph in document.paragraphs
                if paragraph.style.name == "List Number"
            )
        )
        bold_runs = [
            run
            for paragraph in document.paragraphs
            for run in paragraph.runs
            if run.text == "중요:"
        ]
        self.assertEqual(len(bold_runs), 1)
        self.assertTrue(bold_runs[0].bold)
        self.assertTrue(
            any(paragraph.style.name == "Quote" for paragraph in document.paragraphs)
        )
        self.assertEqual(
            sum(
                paragraph.style.name == "List Number"
                for paragraph in document.paragraphs
            ),
            2,
        )

    def test_internal_codex_citation_token_is_rejected_before_docx_generation(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "# 제목\n\n본문 citeturn1search0\n",
                encoding="utf-8",
            )

            with self.assertRaisesRegex(ValueError, "internal Codex citation"):
                generate_docx_report(markdown, output)

        self.assertFalse(output.exists())

    def test_long_cover_title_wraps_at_a_semantic_word_boundary(self) -> None:
        title = "OCI Console AI와 MCP 기반 클라우드 운영 자동화 검토"

        wrapped = wrap_cover_title(title)

        self.assertEqual(
            wrapped,
            "OCI Console AI와 MCP 기반\n클라우드 운영 자동화 검토",
        )
        self.assertNotIn("운\n영", wrapped)

    def test_dynamic_title_and_document_type_replace_hardcoded_minutes_label(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "# MySQL HeatWave 동시성 아키텍처\n\n"
                "문서 유형: 기술 세션 분석\n\n"
                "## 화자별 구성\n\n"
                "- 화자 1: 제품 개요\n"
                "- 화자 2: AutoML 및 GenAI\n\n"
                "## DASK worker 제약\n\n"
                "| 작업 | 제한 |\n"
                "|---|---|\n"
                "| GenAI | 노드당 worker 1개 |\n",
                encoding="utf-8",
            )

            generate_docx_report(
                markdown,
                output,
                document_title="MySQL HeatWave 동시성 아키텍처",
                document_type="기술 세션 분석",
                saved_date="2026-02-12",
            )

            document = Document(output)
            paragraph_text = [paragraph.text for paragraph in document.paragraphs]
            self.assertEqual(paragraph_text.count("MySQL HeatWave 동시성 아키텍처"), 1)
            self.assertIn("기술 세션 분석", paragraph_text)
            self.assertNotIn("회의록", paragraph_text)

            namespace = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }
            attribute = lambda name: f"{{{namespace['w']}}}{name}"
            with ZipFile(output) as archive:
                root_xml = ET.fromstring(archive.read("word/document.xml"))
            anchors = [
                link.get(attribute("anchor"))
                for link in root_xml.findall(".//w:hyperlink", namespace)
            ]
            bookmarks = {
                item.get(attribute("name"))
                for item in root_xml.findall(".//w:bookmarkStart", namespace)
            }
            self.assertTrue(anchors)
            self.assertTrue(all(anchor in bookmarks for anchor in anchors))

    def test_content_headings_use_matching_multilevel_numbers_in_toc_and_body(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "# 기술 세션 분석\n\n"
                "문서 유형: 기술 발표\n\n"
                "## 지원 종료 정책\n\n"
                "### MySQL 8.0 일정\n\n"
                "본문\n\n"
                "## 동시성 아키텍처\n\n"
                "### DASK worker\n\n"
                "본문\n",
                encoding="utf-8",
            )

            generate_docx_report(
                markdown,
                output,
                document_title="기술 세션 분석",
                document_type="기술 발표",
                saved_date="2026-02-12",
            )

            namespace = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }
            with ZipFile(output) as archive:
                root_xml = ET.fromstring(archive.read("word/document.xml"))
            toc_texts = [
                "".join(link.itertext())
                for link in root_xml.findall(".//w:hyperlink", namespace)
            ]
            self.assertEqual(
                toc_texts,
                [
                    "1. 지원 종료 정책",
                    "1.1 MySQL 8.0 일정",
                    "2. 동시성 아키텍처",
                    "2.1 DASK worker",
                ],
            )

            document = Document(output)
            numbered_headings = [
                paragraph
                for paragraph in document.paragraphs
                if paragraph.style.name in {"Heading 1", "Heading 2"}
                and paragraph.text != "목차"
            ]
            self.assertEqual(
                [paragraph.text for paragraph in numbered_headings],
                [
                    "지원 종료 정책",
                    "MySQL 8.0 일정",
                    "동시성 아키텍처",
                    "DASK worker",
                ],
            )
            self.assertEqual(
                [
                    int(
                        paragraph._p.pPr.numPr.ilvl.get(qn("w:val"))
                    )
                    for paragraph in numbered_headings
                ],
                [0, 1, 0, 1],
            )

    def test_dense_toc_collapses_to_top_level_headings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            sections = "\n\n".join(
                f"## Topic {index}\n\n### Detail {index}\n\nBody {index}."
                for index in range(1, 14)
            )
            markdown.write_text(
                "# Dynamic analysis\n\n"
                "Document type: Technical analysis\n\n"
                f"{sections}\n",
                encoding="utf-8",
            )

            generate_docx_report(markdown, output, saved_date="2026-07-15")

            namespace = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }
            with ZipFile(output) as archive:
                root_xml = ET.fromstring(archive.read("word/document.xml"))
            toc_texts = [
                "".join(link.itertext())
                for link in root_xml.findall(".//w:hyperlink", namespace)
            ]

        self.assertEqual(len(toc_texts), 13)
        self.assertTrue(all("Detail" not in text for text in toc_texts))
        self.assertEqual(toc_texts[0], "1. Topic 1")
        self.assertEqual(toc_texts[-1], "13. Topic 13")

    def test_markdown_checklists_render_as_real_checkbox_glyphs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "# Technical analysis\n\n"
                "Document type: Operational brief\n\n"
                "## Actions\n\n"
                "- [ ] Validate the rollout.\n"
                "- [x] Preserve the evidence.\n",
                encoding="utf-8",
            )

            generate_docx_report(markdown, output, saved_date="2026-07-15")
            paragraph_text = [
                paragraph.text for paragraph in Document(output).paragraphs
            ]

        self.assertIn("☐ Validate the rollout.", paragraph_text)
        self.assertIn("☑ Preserve the evidence.", paragraph_text)
        self.assertFalse(any("[ ]" in text or "[x]" in text for text in paragraph_text))

    def test_inline_code_and_official_links_render_without_markdown_syntax(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown = root / "analysis.md"
            output = root / "analysis.docx"
            markdown.write_text(
                "# 기술 세션 분석\n\n"
                "문서 유형: 기술 발표\n\n"
                "## 외부 근거 확인\n\n"
                "- `CALL sys.ML_RAG()` 예시는 "
                "[Oracle 공식 문서](https://docs.example.com/heatwave)에서 확인했다.\n",
                encoding="utf-8",
            )

            generate_docx_report(
                markdown,
                output,
                document_title="기술 세션 분석",
                document_type="기술 발표",
                saved_date="2026-07-13",
            )

            namespace = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
            }
            with ZipFile(output) as archive:
                document_xml = ET.fromstring(archive.read("word/document.xml"))
                relationships_xml = ET.fromstring(
                    archive.read("word/_rels/document.xml.rels")
                )

            document_text = "".join(document_xml.itertext())
            self.assertIn("CALL sys.ML_RAG()", document_text)
            self.assertIn("Oracle 공식 문서", document_text)
            self.assertNotIn("[Oracle 공식 문서]", document_text)
            self.assertNotIn("https://docs.example.com/heatwave", document_text)

            external_links = document_xml.findall(".//w:hyperlink[@r:id]", namespace)
            self.assertEqual(len(external_links), 1)
            targets = {
                relationship.get("Target")
                for relationship in relationships_xml.findall(
                    "pr:Relationship", namespace
                )
                if relationship.get("Type", "").endswith("/hyperlink")
            }
            self.assertEqual(targets, {"https://docs.example.com/heatwave"})


if __name__ == "__main__":
    unittest.main()
